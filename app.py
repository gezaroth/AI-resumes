from docx import Document
import streamlit as st
from typing import Union, List
import os
import time
from PyPDF2 import PdfReader
from langchain.prompts import PromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangDocument
from langchain.chains.qa_with_sources import load_qa_with_sources_chain
from langchain.vectorstores import VectorStore
from langchain_community.vectorstores import Qdrant
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv

load_dotenv()

OPEN_AI_API = st.secrets["OPEN_AI_API_KEY"]
os.environ["OPEN_AI_API_KEY"] = OPEN_AI_API


title_formats = {
    "Título del Documento",
    "Subtitulo del Documento",
    "Cabecera del TDC",
    "Heading 1",
    "Heading 2",
}

template = """You are an expert in summarizing bank documents.
Your task is to analyze each paragraph in detail and summarize it without losing any information.
IMPORTANT: Do not translate or change the language of the text. Keep the response in the original language of the text.
If the text is in English, leave it in English. If the text is in Spanish, leave it in Spanish.
QUESTION: {question}
===========
{summaries}
===========
FINAL ANSWER: """

STUFF_PROMPT = PromptTemplate(
    template=template, input_variables=["summaries", "question"]
)



def text_to_docs(text:Union[str, List[str]]) -> List[LangDocument]:
    """Converts a string or list of strings to a list of Documents
    with metadata."""
    if isinstance(text, str):
        text = [text]
    page_docs = [LangDocument(page_content=page) for page in text]

    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1
    doc_chunks = []
    text_splitter = RecursiveCharacterTextSplitter(
    separators=["\n\n", "\n", " ", ".", ";", ":", ",", "!", "?", "%"],
    )
    for doc in page_docs:
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            chunk_doc = LangDocument(
                page_content=chunk, metadata={"page":doc.metadata["page"], "chunk":i}
            )
            # Add sources a metadata
            chunk_doc.metadata["source"] = f"{chunk_doc.metadata['page']}-{chunk_doc.metadata['chunk']}"
            doc_chunks.append(chunk_doc)
    return doc_chunks

def embeded_docs(docs:List[LangDocument]) -> VectorStore:
    """Embeds a list of Documents and returns a FAISS index"""
    # Embeded chunks
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
    qdrant = Qdrant.from_documents(docs, embeddings, location=":memory:")
    return qdrant

def summarize_text(text: str) -> str:
    text_docs = text_to_docs(text)
    index = embeded_docs(text_docs)
    prompt = f"""You are an ai specialist in legal, financial and technical language.
    The user will give you a technical text and you will have to summarise it.
    Follow these rules:
    - The output text should contain as little text as possible, without losing meaning.
    - Do not mention that you are an AI or anything like that.
    - Summarize the text in the same language as it was originally provided.
    - Do not give any response in another language other than the originally provided.
    - Do not translate to English.
    - If the text is in English, leave it in English.
    The text is the next: {text}
    """
    while True:
        try:
            llm = ChatOpenAI(api_key=OPEN_AI_API, temperature=0.3)
            chain = load_qa_with_sources_chain(llm=llm, chain_type="stuff", prompt=STUFF_PROMPT)
            docs = index.similarity_search(query=prompt, k=1)
            response = chain.run(input_documents=docs, question=prompt)
            return response
        except Exception as e:
            print(f"ERROR: {e}")
            time.sleep(10)

def process_docx(doc_obj):
    documento = Document(doc_obj)
    for parrafo in documento.paragraphs:
        texto = parrafo.text
        if parrafo.style.name in title_formats or len(parrafo.text) < 5:
            continue
        response = summarize_text(texto)
        st.markdown("### PARRAFO ORIGINAL")
        st.markdown(texto)
        st.markdown("### PARRAFO REESCRITO")
        st.markdown("\n")
        st.markdown(response)
        st.markdown("### ----------------------------------------------------------")


def process_pdf(doc_obj):
    pdf_reader = PdfReader(doc_obj)
    for page in pdf_reader.pages:
        texto = page.extract_text()
        response = summarize_text(texto)
        # st.markdown("### PARRAFO ORIGINAL")
        # st.markdown(texto)
        # st.markdown("### PARRAFO REESCRITO")
        # st.markdown("\n")
        st.markdown(response)
        # st.markdown("### ----------------------------------------------------------")

def process_txt(doc_obj):
    texto = doc_obj.read().decode("utf-8")
    response = summarize_text(texto)
    st.markdown("### PARRAFO ORIGINAL")
    st.markdown(texto)
    st.markdown("### PARRAFO REESCRITO")
    st.markdown("\n")
    st.markdown(response)
    st.markdown("### ----------------------------------------------------------")


def main():
    st.header("Document Helper")
    
    allowed_types = ["pdf", "docx", "txt"]
    doc_obj = st.file_uploader("Carga tu documento", type=allowed_types)

    if doc_obj:
        file_type = doc_obj.type
        try:
            if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                process_docx(doc_obj)
            elif file_type == "application/pdf":
                process_pdf(doc_obj)
            elif file_type == "text/plain":
                process_txt(doc_obj)
            else:
                st.warning("Tipo de archivo no compatible. Por favor carga un archivo DOCX, PDF o TXT.")
        except Exception as e:
            st.error(f"Error al procesar el documento: {e}")

if __name__ == "__main__":
    main()